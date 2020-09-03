from pytube import YouTube
import os
import moviepy.editor as mp
import urllib.request
from bs4 import BeautifulSoup
from pydub import AudioSegment
import speech_recognition as sr
import re
from gensim.summarization.summarizer import summarize
from konlpy.tag import Kkma
from konlpy.tag import Okt
from wordcloud import WordCloud
from collections import Counter
import kss
import glob
import docx

class Summerize(object):

    def __init__(self):
        self.kkma = Kkma()
        self.twitter = Okt()
        self.number = 0
        path = os.getcwd() + '\\HMKMRHD.TTF'
        self.wc = WordCloud(font_path=path, background_color="White", width=1000, height=1000, max_words=100, max_font_size=300)


    def makeSentence(self, originaltext):
        #받아온 text 문장으로 나누기

        textList = []
        text = ''

        for sent in kss.split_sentences(originaltext):
            textList.append(sent)

        for i in textList:
            text = text + i + '. '#문장으로 구분하기 위해 '. '추가

        return text


    def text2sentences(self, text):
        #text 요약하기

        #print(summarize(text, ratio=0.3))
        sentences1 = summarize(text, ratio=0.1)

        return self.kkma.sentences(sentences1)


    def get_nouns(self, sentences):
        #문장의 특정 단어 추출

        nouns = []

        for sentence in sentences:
            if sentence != '':
                nouns.append(' '.join([noun for noun in self.twitter.nouns(str(sentence))
                                       if len(noun) > 1]))#길이가 1인 단어는 제외

        return nouns


    def keyword_nouns(self, text, title):
        #키워드 wordcloud 형식으로 변환해주기
        noun = self.twitter.nouns(text)

        for i,v in enumerate(noun):
            if(len(v)<2) :
                noun.pop(i)#글자수 1자 시 제외

        count = Counter(noun)
        noun_list = count.most_common(100)
        result = [i for i in noun_list if i[1] > 1]#빈도수 1회시 제외
        print(result)

        self.wc.generate_from_frequencies(dict(result))
        self.wc.to_file("wordCloud.png")

    def make_file(self, sentence1, sentence2, title):
        # 요약 문장, 단어 파일에 담기

        doc = docx.Document()
        doc.add_heading(title+"영상 요약", 0)
        para = doc.add_heading("<줄거리 요약>\n", 1)

        for sentence in sentence1:
            print(sentence)
            para = doc.add_paragraph()
            run = para.add_run(sentence)

        para = doc.add_heading("\n\n<문장 키워드 요약>\n", 1)

        for sentence in sentence2:
            print(sentence)
            para = doc.add_paragraph()
            run = para.add_run(sentence)

        para = doc.add_heading("\n\n<영상 Keyword>\n", 1)
        doc.add_picture(os.getcwd() + '\\wordCloud.png', width=docx.shared.Cm(10), height=docx.shared.Cm(10))  # 위치 바꾸기
        doc.save(title+'.docx')  # 저장 이름 바꾸기


class AudioToText(object):
    def __init__(self, path, title, num):
        self.path = path+title+str(num)+".wav"

    def makeText(self): #음성 파일 텍스트로 변환
        r = sr.Recognizer()
        file = sr.AudioFile(self.path)
        with file as source:
            audio = r.record(source)
        return r.recognize_google(audio, language='ko-KR')

class YouTubeDownloader(object): #youtube에서 동영상 다운 받기
    def __init__(self, youtube_rul):
        self.path = './movie'
        self.youtube_url = youtube_rul

    def downloader(self):
        yt = YouTube(self.youtube_url)
        os.makedirs(self.path, exist_ok=True) #영상 파일을 저장할 하위 디렉터리 생성
        try:
            source = urllib.request.urlopen(self.youtube_url).read()
        except ValueError as er:
            print("잘못된 url 입니다")
            return 0
        soup = BeautifulSoup(source, "html.parser")
        title = soup.find('title').text #페이지 크롤링하여 영상 제목 가져오기
        title = re.sub("[-=+,#/\?:^$@*\"※~&%ㆍ!』\\‘|\(\)\[\]\<\>`\'》]", "", title) #영상 제목에서 특수 문자 제외하기
        print(self.path)
        yt.streams.filter(progressive=True, file_extension='mp4').order_by('resolution').desc().first().download(self.path) #영상 다운 받기
        if os.path.isfile(self.path+'Unknown YouTube Video Title.mp4'):
            os.rename(self.path+'Unknown YouTube Video Title.mp4', self.path+title+'.mp4') #영상 제목 바꾸기
        else:
            other_file = glob.glob(self.path+title[0:3]+"*.mp4")
            other_title = os.path.basename(other_file[0])
            os.rename(self.path+other_title, self.path+title+".mp4")
        return title

    def pathvalue(self):
        self.path += '/'
        return self.path

class ChangeMovie(object):
    def __init__(self, path, filename):
        self.path = path
        self.filename = filename

    def MovieToAudio(self):
        clip = mp.VideoFileClip(self.path+self.filename+".mp4")
        clip.audio.write_audiofile(self.path + self.filename + ".mp3") #영상을 음원 파일로 바꾸기

        filedir = str(os.getcwd() + "/movie/" + self.filename)
        src = filedir + ".mp3"
        print(src)
        sound = AudioSegment.from_mp3(src)
        num = 0
        movie_length = (len(sound)//1000)*1000
        for i in range(0, movie_length, 180000): #3분 단위로 음원 파일 잘라서 형식 바꾸기
            if i+180000 > movie_length:
                temp = sound[i:]
            else:
                temp = sound[i:i+180000]
            temp.export(filedir+str(i)+".wav", format="wav")
            num += 1
        return num

url = input("url 입력: ")
y = YouTubeDownloader(url)
yp = y.pathvalue()
yt = y.downloader()
original_text = ""
if yt != 0:
    c = ChangeMovie(yp, yt)
    num = c.MovieToAudio()
    for i in range(0, num):
        a = AudioToText(yp, yt, i*180000)
        original_text += str(a.makeText())
        original_text += " "
    print(original_text)

#순서 : originaltext 받아와 문장 구분을 위해 makeSentence 함수 호출 -> 문장이 구분된 text 요약하기 위해 text2sentence 함수 호출 -> 파일에 쓰기 위해 make_file함수 호출 -> 추가적으로 중요 명사들 wordcloud 하기 위해 keyword_nouns함수 호출
summer = Summerize()
text = summer.makeSentence(original_text)
summer.keyword_nouns(text, yt)

complete_sentences1 = summer.text2sentences(text)
complete_sentences2 = summer.get_nouns(summer.text2sentences(text))

summer.make_file(complete_sentences1, complete_sentences2, yt)